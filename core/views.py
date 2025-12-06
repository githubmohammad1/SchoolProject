# core/views.py 
from rest_framework import generics, status, viewsets
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from .permissions import *
from rest_framework import permissions
from django.shortcuts import render
from django.http import HttpResponse

from docx import Document
from io import BytesIO
# الاستيراد الموحد لجميع النماذج (باستخدام * مناسب هنا)
from .models import * # الاستيراد الموحد لجميع الـ Serializers 
from .serializers import *
from docx.enum.text import WD_ALIGN_PARAGRAPH # للاستخدام في محاذاة النص
from docx.shared import Inches # لتحديد عرض الأعمدة

def export_student_list_docx(request):
    """
    تجميع بيانات الطلاب وتصديرها كملف Word (.docx) مع دعم RTL والتوقيعات.
    """
    # 🛡️ التحقق من الصلاحيات (معلم أو مدير فقط)
    if not request.user.is_staff and not (hasattr(request.user, 'teacher') and request.user.teacher):
         return HttpResponse("غير مصرح لك باستخراج هذا التقرير.", status=403)
         
    # 📚 جلب بيانات الطلاب (مع جلب البيانات المرتبطة لتحسين الأداء)
    students = Student.objects.select_related('user', 'class_ref', 'parent__user').all().order_by('class_ref__name', 'user__username')
    
    # 1. إنشاء مستند Word جديد
    document = Document()
    
    # 2. إضافة عنوان (رئيسي) ومحاذاته لليمين (RTL)
    heading = document.add_heading('قائمة الطلاب الرسمية في المدرسة', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT # محاذاة عنوان التقرير لليمين

    # 3. إنشاء جدول وإضافة الرؤوس
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'الرقم'
    hdr_cells[1].text = 'اسم الطالب'
    hdr_cells[2].text = 'الصف'
    hdr_cells[3].text = 'تاريخ الميلاد'
    hdr_cells[4].text = 'ولي الأمر'
    
    # تطبيق محاذاة لليمين على رؤوس الأعمدة
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 4. ملء صفوف الجدول ببيانات الطلاب
    for i, student_profile in enumerate(students):
        row_cells = table.add_row().cells
        
        # تجميع أسماء الطلاب والأولياء
        student_name = f"{student_profile.user.first_name} {student_profile.user.last_name}" if student_profile.user.first_name else student_profile.user.username
        parent_name = student_profile.parent.user.username if student_profile.parent and student_profile.parent.user else 'لا يوجد'
        
        # تصحيح الخطأ: الوصول إلى تاريخ الميلاد عبر student_profile.user
        dob = student_profile.date_of_birth.strftime('%Y-%m-%d') if student_profile.date_of_birth else 'غير محدد'

        row_cells[0].text = str(i + 1)
        row_cells[1].text = student_name
        row_cells[2].text = student_profile.class_ref.name if student_profile.class_ref else 'غير محدد'
        row_cells[3].text = dob
        row_cells[4].text = parent_name
        
        # تطبيق محاذاة لليمين على محتوى الخلايا
        for cell in row_cells:
             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # -----------------------------------------------------
    # 5. إضافة قسم التوقيعات (Signatures Section)
    # -----------------------------------------------------
    
    # إضافة مسافة فاصلة بعد الجدول
    document.add_paragraph()
    document.add_paragraph() 
    
    # إنشاء جدول بسيط بصف واحد وعمودين لضمان تنسيق التوقيعات
    signature_table = document.add_table(rows=1, cols=2)
    
    # تعيين عرض الأعمدة لتقسيم الصفحة
    signature_table.columns[0].width = Inches(3.0) # أمين السر
    signature_table.columns[1].width = Inches(3.0) # المدير

    # الخلية اليمنى (المدير)
    manager_cell = signature_table.cell(0, 1)
    manager_paragraph = manager_cell.paragraphs[0]
    manager_paragraph.text = "المدير: ________________"
    # دفع توقيع المدير إلى أقصى اليمين
    manager_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # الخلية اليسرى (أمين السر)
    secretary_cell = signature_table.cell(0, 0)
    secretary_paragraph = secretary_cell.paragraphs[0]
    secretary_paragraph.text = "أمين السر: _____________"
    # دفع توقيع أمين السر إلى أقصى اليسار
    secretary_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 6. حفظ المستند وإرجاع الاستجابة (باقي الكود كما هو)
    f = BytesIO()
    document.save(f)
    f.seek(0)
    
    filename = 'Student_List.docx'
    response = HttpResponse(
        f.read(), 
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response
# ************
class SchoolInfoViewSet(viewsets.ModelViewSet):
    """
    إدارة معلومات المدرسة (المدير، أمين السر، إلخ). 
    يجب أن يكون الوصول مقيدًا بالإدارة العليا فقط.
    """
    queryset = SchoolInfo.objects.all()
    serializer_class = SchoolInfoSerializer
    # TODO: إضافة صلاحيات مخصصة (مثل IsPrincipalOrAdmin)
    permission_classes = [IsSchoolAdministrator | IsReadOnly] 

class AssessmentTypeViewSet(viewsets.ModelViewSet):
    """
    إدارة أنواع التقييمات وأوزانها (مذاكرة، مشروع، امتحان نهائي).
    """
    queryset = AssessmentType.objects.all()
    serializer_class = AssessmentTypeSerializer
    # TODO: إضافة صلاحيات مخصصة (مثل IsAdminUser)
    # الصلاحية: المدرسون والمشرفون فقط هم من يمكنهم تعريف أنواع التقييمات
    permission_classes = [IsTeacherOrAdmin | IsReadOnly]

class BehaviorTypeViewSet(viewsets.ModelViewSet):
    """
    إدارة أنواع السلوكيات الإيجابية والسلبية وقيم النقاط.
    """
    queryset = BehaviorType.objects.all()
    serializer_class = BehaviorTypeSerializer
    # TODO: إضافة صلاحيات مخصصة (مثل IsAdminUser)
    permission_classes = [IsTeacherOrAdmin | IsReadOnly]


class BehaviorRecordViewSet(viewsets.ModelViewSet):
    """
    تسجيل وعرض السلوكيات (نقاط السلوك).
    يجب أن يتمكن المدرسون من التسجيل، والإدارة من العرض والتحرير.
    """
    # يمكن للمدرسين رؤية كل السجلات التي تخص طلابهم أو جميع السجلات إذا كانوا إداريين
    queryset = BehaviorRecord.objects.all() 
    serializer_class = BehaviorRecordSerializer
    # TODO: إضافة صلاحيات مخصصة لضمان أن المدرسين فقط هم من يمكنهم الإنشاء
   # الصلاحية: القراءة مسموحة للجميع، والكتابة/الإنشاء فقط للمعلمين
    permission_classes = [IsTeacherOrGuidance | IsReadOnly]
    
    def perform_create(self, serializer):
        """تحديد المدرس الذي قام بتسجيل السلوك تلقائيًا."""
        user = self.request.user
        
        # التأكد من أن المستخدم الحالي هو مدرس
        if not hasattr(user, 'teacher'):
            # استخدام PermissionDenied إذا لم يكن المستخدم معلمًا
            raise exceptions.PermissionDenied("Only teachers are allowed to record behavior.")

        # حفظ السجل وتعيين المدرس الحالي كـ recorded_by
        serializer.save(recorded_by=user.teacher)
class GradeViewSet(viewsets.ModelViewSet):
    # ندمج الصلاحيتين: إما أن تكون مدرساً/مشرفاً، أو مجرد مستخدم مسجل للدخول (للقراءة فقط)
    permission_classes = [IsTeacherOrAdmin | permissions.IsAuthenticated]
    serializer_class = GradeSerializer

    def get_queryset(self):
        user = self.request.user
        # للمدرسين والمشرفين: عرض جميع الدرجات أو درجات دوراته
        if user.is_staff or hasattr(user, 'teacher'):
            if user.is_staff:
                 return Grade.objects.all()
            # المدرس يرى درجات طلابه في الدورات التي يدرسها
            teacher_courses = user.teacher.teaching_courses.all()
            return Grade.objects.filter(assignment__course__in=teacher_courses)

        # للطالب: عرض درجاته فقط
        elif hasattr(user, 'student'):
            return Grade.objects.filter(student=user.student)

        # لولي الأمر: عرض درجات أبنائه فقط
        elif hasattr(user, 'parentprofile'):
            children_students = user.parentprofile.children.all()
            return Grade.objects.filter(student__in=children_students)
        
        # لغير ذلك (مجرد مستخدم مسجل): لا يرى شيئاً
        return Grade.objects.none()

class ClassViewSet(viewsets.ModelViewSet):
    queryset = Class.objects.all()
    serializer_class = ClassSerializer
    permission_classes = [IsAuthenticated, IsTeacherOrAdmin]



class SubjectViewSet(viewsets.ModelViewSet):
    queryset = Subject.objects.all()
    serializer_class = SubjectSerializer
    permission_classes = [IsAuthenticated, IsTeacherOrAdmin]



class CourseViewSet(viewsets.ModelViewSet):
    """
    إدارة المواد الدراسية.
    - المعلم يرى المواد التي يدرسها.
    - الطالب يرى المواد الخاصة بصفه.
    - المشرف يرى كل شيء.
    """
    serializer_class = CourseSerializer
    permission_classes = [permissions.IsAuthenticated] # أي مستخدم مسجل يمكنه القراءة

    def get_queryset(self):
        user = self.request.user

        # 1. إذا كان المشرف أو المدير
        if user.is_staff or user.is_superuser:
            return Course.objects.all()

        # 2. إذا كان معلماً
        if hasattr(user, 'teacher'):
            # يرجع المواد التي تم ربط هذا المعلم بها تحديداً
            return Course.objects.filter(teacher=user.teacher)

        # 3. إذا كان طالباً
        if hasattr(user, 'student'):
            student_profile = user.student
            # نتأكد أولاً أن الطالب مربوط بصف (class_ref)
            if student_profile.class_ref:
                # يرجع كل المواد المرتبطة بنفس صف الطالب
                return Course.objects.filter(class_level=student_profile.class_ref)
            else:
                return Course.objects.none() # طالب بلا صف لا يرى مواد

        # 4. إذا كان ولي أمر (اختياري: يرى مواد أبنائه)
        if hasattr(user, 'parentprofile'):
             # يمكننا تركها فارغة حالياً أو جلب مواد الأبناء
             return Course.objects.none()

        return Course.objects.none()
class StudentViewSet(viewsets.ModelViewSet):
    serializer_class = StudentProfileSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        # المعلم يرى جميع الطلاب (أو يمكن فلترتهم حسب الصفوف التي يدرسها)
        if hasattr(user, 'teacher') or user.is_staff:
            return Student.objects.all()
        return Student.objects.none()
class RegisterView(generics.CreateAPIView):
    serializer_class = UserRegisterSerializer
    # يجب أن يكون مسموحاً للجميع بالتسجيل
    permission_classes = () 

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        user = serializer.save()

        return Response({
            "user": serializer.data,
            "message": "تم إنشاء الحساب بنجاح. يمكنك الآن تسجيل الدخول."
        }, status=status.HTTP_201_CREATED)
# 2. عرض الملف الشخصي (Profile View)
# ----------------------------------------
class ProfileView(APIView):
    permission_classes = (IsAuthenticated,) 

    def get(self, request, *args, **kwargs):
        user = request.user
        profile_data = {}
        role = 'guest' # القيمة الافتراضية

        # 1. فحص المدير (Admin)
        if user.is_staff or user.is_superuser:
            role = 'admin'
            profile_data = UserSerializer(user).data
            profile_data['id'] = user.id

        # 2. فحص المعلم (Teacher)
        # نستخدم try/except لأنها أدق من hasattr في بعض الحالات
        elif self._is_teacher(user):
            role = 'teacher'
            profile_data = TeacherProfileSerializer(user.teacher).data

        # 3. فحص الطالب (Student)
        elif self._is_student(user):
            role = 'student'
            profile_data = StudentProfileSerializer(user.student).data

        # 4. فحص ولي الأمر (Parent)
        elif self._is_parent(user):
            role = 'parent'
            profile_data = ParentProfileSerializer(user.parentprofile).data
            
        else:
            return Response({
                "message": "المستخدم ليس له دور محدد (معلم/طالب/ولي أمر)",
                "role": "guest"
            }, status=status.HTTP_200_OK) # نرجع 200 بدل 404 لكي لا ينهار التطبيق
        
        return Response({
            "status": "success",
            "role": role,
            "profile": profile_data
        }, status=status.HTTP_200_OK)

    # دوال مساعدة للتحقق (Helpers)
    def _is_teacher(self, user):
        try:
            return user.teacher is not None
        except Teacher.DoesNotExist:
            return False

    def _is_student(self, user):
        try:
            return user.student is not None
        except Student.DoesNotExist:
            return False

    def _is_parent(self, user):
        try:
            return user.parentprofile is not None
        except ParentProfile.DoesNotExist:
            return False


class NotificationViewSet(viewsets.ModelViewSet):
    serializer_class = NotificationSerializer
    # الإشعارات للقراءة فقط للمستخدمين العاديين، ويتم إنشاؤها عبر المشرف/المدرس/النظام
    permission_classes = [permissions.IsAuthenticated] 

    def get_queryset(self):
        # كل مستخدم يرى إشعاراته فقط
        return Notification.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        # للتأكد من ربط الإشعار بالمستخدم أثناء الإنشاء عبر API
        # (يمكن لاحقاً منع الإنشاء إلا للمشرفين إذا كانت الإشعارات تولد فقط من النظام)
        if self.request.user.is_staff or hasattr(self.request.user, 'teacher'):
            serializer.save()
        # نتركها فارغة هنا، حيث سيتم توليد الإشعارات عبر الـ signals وليس عبر واجهة API مباشرة في معظم الحالات.

class AssignmentViewSet(viewsets.ModelViewSet):
    queryset = Assignment.objects.all()
    # نسمح بالقراءة للمستخدمين العاديين، والتعديل/الإضافة للمدرس/المشرف
    permission_classes = [IsTeacherOrAdmin | permissions.IsAuthenticated]
    serializer_class = AssignmentSerializer

    def get_queryset(self):
        user = self.request.user
        
        # المشرف يرى جميع الواجبات
        if user.is_staff:
            return Assignment.objects.all()
            
        # المدرس يرى واجبات دوراته
        if hasattr(user, 'teacher'):
            teacher_courses = user.teacher.teaching_courses.all()
            return Assignment.objects.filter(course__in=teacher_courses)

        # الطالب يرى الواجبات لدورات صفه
        if hasattr(user, 'student'):
            student_class = user.student.class_ref
            return Assignment.objects.filter(course__class_level=student_class)
            
        return Assignment.objects.none()